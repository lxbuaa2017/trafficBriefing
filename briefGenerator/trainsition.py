# -*- coding: utf-8 -*-
import datetime
import json

import docx
import jieba
import re
import jieba
import time
import os
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches
from docx.oxml.ns import qn

p=os.path.dirname(__file__)
p = os.path.join(p,os.path.pardir)
raw = {}
with open('{}/trafficBriefing/trafficBriefing/brief.json'.format(p), 'r', encoding='utf-8') as fObj:
    raw_list = json.load(fObj)
    raw = raw_list[-1]


path='{}/brief.docx'.format(p)
if os.path.exists(path):  # 如果文件存在
    os.remove(path)  
fd = open(path, mode="w", encoding="utf-8")
fd.close()
document=Document()
now = datetime.datetime.now()
img_url = "{}/trafficBriefing/trafficBriefing/img/".format(p)+str(now.year) + '年' + str(now.month) + '月' + str(now.day) + '日' + "全国交通气象预报"+".png"
document.add_picture(img_url, width=Inches(6.0748031))


strs = "陆域。"
content = ""


lenth = len(raw['snow_influenced'])
if lenth>1:
    strs+=raw['snow_influenced'][0][:-1]
    strs+=raw['snow_influenced'][1]
    strs+="等" + str(lenth - 1) + "条"


lenth = len(raw['fog_influenced'])
if lenth>1:
    strs+='，'
    strs+=raw['fog_influenced'][0][:-1]
    strs+=raw['fog_influenced'][1]
    strs+="等" + str(lenth - 1) + "条"


lenth = len(raw['rain_influenced'])
if lenth>1:
    strs+='，'
    strs+=raw['rain_influenced'][0][:-1]
    strs+=raw['rain_influenced'][1]
    strs+="等" + str(lenth - 1) + "条"


lenth = len(raw['thunder_influenced'])
if lenth>1:
    strs+='，'
    strs+=raw['thunder_influenced'][0][:-1]
    strs+=raw['thunder_influenced'][1]
    strs+="等" + str(lenth - 1) + "条"


lenth = len(raw['other_influenced'])
if lenth>0:
    for influenced in raw['other_influenced']:
        strs+='，'
        strs+=influenced[0][:-1]
        strs+=influenced[1]
        strs+="等" + str(len(influenced) - 1) + "条"


strs+='。'



p1 = document.add_paragraph().add_run(strs)
p1.font.name = u'仿宋_GB2312'
p1._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
# 小四对应12pt
p1.font.size = Pt(12)
p1.add_break(docx.enum.text.WD_BREAK.PAGE)



p2 = document.add_paragraph().add_run("附件")
p2.font.name = u'方正小标宋简体'
p2._element.rPr.rFonts.set(qn('w:eastAsia'), u'方正小标宋简体')
# 小四对应12pt
p2.font.size = Pt(14)
title = "受天气影响的主要路段、海域、机场"

document.save(path)




